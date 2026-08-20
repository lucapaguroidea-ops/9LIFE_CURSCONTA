"""Markdown → .docx, pentru documentele revizuite.

De ce un convertor propriu: `pandoc` nu e disponibil aici, iar LibreOffice headless nu
convertește în acest mediu. `.docx`-urile tale existente vin din alt lanț (încorporează
fonturi RobotoMono) și NU pot fi reproduse identic — de aceea rezultatul de aici merge în
`dist/`, iar cele din `surse/` rămân neatinse.

Ce acoperă: exact ce folosesc documentele — titluri, paragrafe, tabele, blocuri de cod,
liste, citate, linii de separare și marcajele inline (bold, italic, cod). Nu e un
convertor Markdown general și nu pretinde să fie.
"""
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

BLEUMARIN = RGBColor(0x1F, 0x4E, 0x79)
GRI = RGBColor(0x5D, 0x6E, 0x7E)

RE_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")


def _stiluri(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(6)
    for nume, marime in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 11.5)):
        st = doc.styles[nume]
        st.font.name = "Calibri"
        st.font.size = Pt(marime)
        st.font.color.rgb = BLEUMARIN
        st.font.bold = True


def _inline(par, text):
    """Scrie text cu marcajele **bold**, *italic* și `cod`."""
    for bucata in RE_INLINE.split(text):
        if not bucata:
            continue
        if bucata.startswith("**") and bucata.endswith("**"):
            par.add_run(bucata[2:-2]).bold = True
        elif bucata.startswith("`") and bucata.endswith("`"):
            r = par.add_run(bucata[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = BLEUMARIN
        elif bucata.startswith("*") and bucata.endswith("*") and len(bucata) > 2:
            par.add_run(bucata[1:-1]).italic = True
        else:
            par.add_run(bucata)


def _cuprins(doc):
    """Câmp de cuprins — Word îl populează la deschidere (F9 dacă nu o face singur)."""
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    run = OxmlElement("w:r")
    txt = OxmlElement("w:t")
    txt.text = "Cuprinsul se generează la deschiderea în Word (Ctrl+A, apoi F9)."
    run.append(txt)
    fld.append(run)
    p._p.append(fld)


def _tabel(doc, randuri):
    """Un tabel Markdown → tabel Word. Prima linie e antet, a doua e separatorul."""
    celule = [[c.strip() for c in r.strip().strip("|").split("|")] for r in randuri]
    celule = [c for i, c in enumerate(celule) if i != 1]      # scoate |---|---|
    lat = max(len(r) for r in celule)
    t = doc.add_table(rows=0, cols=lat)
    t.style = "Light Grid Accent 1"
    for i, rand in enumerate(celule):
        cells = t.add_row().cells
        for j in range(lat):
            cells[j].text = ""
            par = cells[j].paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            _inline(par, rand[j] if j < len(rand) else "")
            for run in par.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
    doc.add_paragraph()


def converteste(md_text, cale_iesire):
    doc = Document()
    _stiluri(doc)

    linii = md_text.split("\n")
    i = 0
    pus_cuprins = False
    while i < len(linii):
        l = linii[i]
        s = l.strip()

        if s.startswith("```"):
            i += 1
            cod = []
            while i < len(linii) and not linii[i].strip().startswith("```"):
                cod.append(linii[i])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            r = p.add_run("\n".join(cod))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            i += 1
            continue

        if s.startswith("|") and i + 1 < len(linii) and set(linii[i + 1].strip()) <= set("|-: "):
            tab = []
            while i < len(linii) and linii[i].strip().startswith("|"):
                tab.append(linii[i])
                i += 1
            _tabel(doc, tab)
            continue

        if s.startswith("#"):
            nivel = len(s) - len(s.lstrip("#"))
            text = s.lstrip("#").strip()
            doc.add_heading(text, level=min(nivel, 3))
            if nivel == 1 and not pus_cuprins:
                _cuprins(doc)
                pus_cuprins = True
            i += 1
            continue

        if s in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            pr = p._p.get_or_add_pPr()
            bdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:color"), "DBE3EA")
            bdr.append(bot)
            pr.append(bdr)
            i += 1
            continue

        if s.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            _inline(p, s[2:])
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = GRI
            i += 1
            continue

        if re.match(r"^[-*+]\s+", s):
            p = doc.add_paragraph(style="List Bullet")
            _inline(p, re.sub(r"^[-*+]\s+", "", s))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", s):
            p = doc.add_paragraph(style="List Number")
            _inline(p, re.sub(r"^\d+\.\s+", "", s))
            i += 1
            continue

        if not s:
            i += 1
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _inline(p, s)
        i += 1

    doc.save(cale_iesire)
    return cale_iesire
